import openai
import pandas as pd
from datetime import datetime
import matplotlib.pyplot as plt
from langchain.chat_models import ChatOpenAI
from docx import Document
from docx.shared import Inches
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import streamlit as st

def get_core_data_tests_summary(core_validation_output, total_validations, total_failed_counts):
    f"""
    Generates a plain-text summary of the core data validations, 
    including overall checks performed, failed checks, data quality,
    and critical categories—without any heading in the response.

    Parameters:
    - core_validation_output (list): A list of JSON-like dictionaries, each representing
      validation results for a specific category (not used in this function, but included for completeness).
    - total_validations (int): The total number of validation checks performed.
    - total_failed_counts (int): The total number of failed validation checks across all categories.

    Returns:
    - str: A plain-text summary with no heading.
    """

    prompt = f"""
    You are a data validation analyst with expertise in credit risk data analysis. 
    Your task is to generate a concise overview (summary) of validation results 
    across multiple categories of a credit risk datamart.

    **Validation Results Input**:
    {core_validation_output}

    Each element of the input list contains:
    - "dataset_name": The category name (e.g., "Borrower Info").
    - "dataset": An object with:
        - "section": A section identifier.
        - "issues": A list of issues, each containing "check_id", "description", "column", and "failed_count".
    - "summary_df": A string representation of a DataFrame summarizing validation results for that category.

    The validation results are represented by the JSON objects (not shown here in detail), 
    but you only need to produce an overall summary that includes:
      - Total number of validation checks performed: {total_validations}
      - Total number of issues identified: {total_failed_counts}
      - Data quality score or percentage of checks that passed
      - The most critical test categories (by name) and *why* they are critical

    **Formatting**:
    - You should write Summary at the start in the font of a subheader

    Now provide the summary:
    """

    # Example: Using LangChain's ChatOpenAI. Adjust as necessary for your environment.
    llm = ChatOpenAI(
        openai_api_key=st.session_state.OPENAI_API_KEY,
        model_name="gpt-4o"
    )
    response = llm.predict(prompt)

    return response

def get_core_data_tests_category_details(core_validation_output):
    """
    Generates plain-text, category-wise details of validation results,
    including checks performed, issues found, impact, and recommendations—
    without any heading in the response.

    Parameters:
    - core_validation_output (list): A list of JSON-like dictionaries, each representing
      validation results for a specific category.

    Returns:
    - str: A category-by-category details text with no heading.
    """

    prompt = f"""
    You are a data validation analyst with expertise in credit risk data analysis. 
    Your task is to provide detailed, category-wise insights into the validation results. 

    Each category has:
      - 'dataset_name' (category name),
      - 'dataset' (details about section and issues),
      - 'summary_df' (a string representation of the summary DataFrame).

    From the provided information (not shown here in full detail), 
    produce a text that covers for each category in very detail:
      - How many checks were performed.
      - How many issues were identified.
      - The approximate percentage or scale of data affected.
      - The most critical issues in that category and why they matter.
      - Recommendations for improving data quality in that category.

    **Formatting**:
    - You should write the name of each category at the start in the font of a subheader

    Here is the list of category outputs to analyze:
    {core_validation_output}

    Now provide only the category-wise details and recommendations in plain text:
    """

    llm = ChatOpenAI(
        openai_api_key=st.session_state.OPENAI_API_KEY,
        model_name="gpt-4o"
    )
    response = llm.predict(prompt)

    return response

# def summarize_core_data_tests(core_validation_output, total_validations, total_failed_counts):
#     """
#     Summarizes core validation test results for credit risk datamart validation.

#     Parameters:
#     - core_validation_output (list): A list of JSON-like dictionaries, each representing
#       validation results for a specific category.
#     - total_validations (int): The total number of validation checks performed.
#     - total_failed_counts (int): The total number of failed validation checks across all categories.

#     Returns:
#     - str: A JSON-formatted string with exactly two keys, each containing text only:
#         "1. Summary": A textual overview of total checks, fails, data quality score, and critical categories.
#         "2. Category-wise Summaries": A textual breakdown of each category’s results plus recommendations.
#     """

#     prompt = f"""
#     You are a data validation analyst with expertise in credit risk data analysis. Your task is to analyze the validation results across multiple categories of a credit risk datamart for a bank.

#     **Validation Results Input**:
#     {core_validation_output}

#     Each element of the input list contains:
#     - "dataset_name": The category name (e.g., "Borrower Info").
#     - "dataset": An object with:
#         - "section": A section identifier.
#         - "issues": A list of issues, each containing "check_id", "description", "column", and "failed_count".
#     - "summary_df": A string representation of a DataFrame summarizing validation results for that category.

#     **Key Parameters**:
#     - Total validations performed: {total_validations}
#     - Total failed checks across categories: {total_failed_counts}

#     **Analysis Requirements**:
#     1. Provide an overall summary covering:
#        - Total number of validation checks performed
#        - Total number of issues identified
#        - Overall data quality score (percentage of checks that passed)
#        - Most critical categories (and why they are critical)
#     2. Provide category-wise summaries:
#        - Number of checks performed
#        - Number of issues identified
#        - Percentage of data affected
#        - Most critical issues and why they matter
#        - Recommendations for improvement

#     **Output Formatting Instructions**:
#     - You MUST return a valid JSON object with exactly two top-level keys:
#         "1. Summary" and "2. Category-wise Summaries"
#     - Each key’s value must be text only (no nested JSON).
#     - Example structure:
#       {{
#         "1. Summary": "Text describing overall findings...",
#         "2. Category-wise Summaries": "Text describing each category plus recommendations..."
#       }}
#     - Do not include any additional keys, objects, or arrays in your JSON. Only these two keys are allowed.
#     """

#     # Example: using a ChatOpenAI interface from LangChain (adjust as needed for your environment)
#     llm = ChatOpenAI(
#         openai_api_key=st.session_state.OPENAI_API_KEY,
#         model_name="gpt-4o"
#     )
#     response = llm.predict(prompt)

#     return response

def analyze_statistical_test(test_name, test_output, threshold, model_type):
    """
    Analyzes the result of a statistical test and generates a structured report using OpenAI.

    Parameters:
    - test_name (str): Name of the statistical test (e.g., T-test, ANOVA, Chi-square, Regression).
    - test_output (float, dict, or DataFrame): The result of the statistical test (e.g., p-value, test statistic, or a DataFrame containing multiple results).
    - threshold (float, optional, or DataFrame): The threshold for significance (e.g., 0.05 for p-values). If None, the analysis should not consider a threshold.
    - model_type (str): The type of model the test is applied to (PD, LGD, or EAD).

    Returns:
    - str: A detailed LLM-generated analysis of the test results.
    """

#     Construct the LLM prompt
    # prompt = f"""
    # You are a financial risk expert helping to interpret the results of a statistical test.

    # Test Name: {test_name}
    # Model Type: {model_type}
    # Test Output: {test_output}
    # Threshold: {threshold if threshold is not None else "Not provided"}

    # Provide a professional and detailed analysis suitable for inclusion in a formal risk management report.
    # If a threshold is provided, interpret the results relative to the threshold and also describe the thresholds accordingly.
    # If no threshold is provided, analyze the results without referencing a threshold.
    # If the test_output is a DataFrame, analyze the key metrics and trends within the data.
    # Ensure the response explains the test's purpose, interprets the results, and provides relevant insights without explicitly labeling sections as 'Introduction', 'Interpretation', or 'Conclusion'.
    # Do not use bold formatting in your response.
    # """

    prompt = f"""
    You are a financial risk expert tasked with interpreting the results of a statistical test for inclusion in a formal risk management report. Your analysis should be detailed, professional, and tailored to an audience familiar with financial risk concepts.

    **Test Details:**
    - Test Name: {test_name}
    - Model Type: {model_type}
    - Test Output: {test_output}
    - Threshold: {threshold if threshold is not None else "Not provided"}

    **Instructions for Analysis:**
    1. **Purpose of the Test:**
    - Begin by explaining the purpose of the {test_name} and its relevance in the context of financial risk management. Describe how this test is typically used to evaluate risk, assess model performance, or validate assumptions. Also, you are not supposed to provide headers in response.

    2. **Interpretation of Results:**
    - If the test output is a numerical value or statistic:
        - Compare the result to the provided threshold (if available). Explain what the threshold represents (e.g., a confidence level, a critical value, or a regulatory benchmark) and how it is derived.
        - Discuss whether the result exceeds, meets, or falls below the threshold and what this implies for the model or risk assessment. For example, does it indicate a significant risk, a model deficiency, or a satisfactory outcome?
        - If no threshold is provided, analyze the result in absolute terms and explain its significance in the context of the test's purpose.
    - If the test output is a DataFrame or a set of metrics:
        - Identify and analyze the key metrics, trends, or patterns within the data. Highlight any values that are particularly high, low, or anomalous.
        - Discuss the implications of these findings for risk management, model performance, or decision-making.

    3. **Threshold Explanation (if applicable):**
    - If a threshold is provided, explain its significance in detail. For example:
        - Is it a statistical threshold (e.g., p-value, confidence interval)? If so, explain how it relates to hypothesis testing and what it indicates about the validity of the results.
        - Is it a regulatory or operational threshold? If so, describe its origin (e.g., industry standards, internal policies) and its importance in risk management.
        - Provide context on how the threshold was determined and why it is relevant to the test.

    4. **Insights and Recommendations:**
    - Based on the results and their interpretation, provide actionable insights or recommendations. For example:
        - If the results indicate a potential risk, suggest mitigation strategies or further investigation.
        - If the results validate the model or assumptions, discuss the implications for decision-making or risk tolerance.
        - If the results are inconclusive or raise additional questions, propose next steps for analysis or testing.

    5. **General Guidelines:**
    - Ensure the analysis is cohesive and flows logically, avoiding explicit section labels like "Introduction" or "Conclusion."
    - Avoid using bold formatting or overly technical jargon unless necessary.
    - Maintain a professional tone suitable for a formal report.
    """
    
    # Call OpenAI API using LangChain
    llm = ChatOpenAI(openai_api_key=st.session_state.OPENAI_API_KEY, model_name="gpt-4o")
    response = llm.predict(prompt)
    
    return response

def summarize_dataset(dataset):
    """
    Generates a summary of the dataset using LangChain.
    
    Parameters:
        - dataset (pd.DataFrame): The dataset to summarize.
    
    Returns:
        - str: A detailed summary of the dataset.
    """
    prompt = f"""
    You are a data analyst reviewing a dataset for a financial risk model.
    
    Dataset Summary:
    Columns: {list(dataset.columns)}
    Data Types: {dataset.dtypes.to_dict()}
    Missing Values: {dataset.isnull().sum().to_dict()}
    Basic Statistics: {dataset.describe().to_dict()}
    
    Provide a detailed summary of the dataset, including key insights, potential issues, and recommendations for improvement. Do not use headings or bold formatting in your response.
    """
    
    # Call OpenAI API using LangChain
    llm = ChatOpenAI(openai_api_key=st.session_state.OPENAI_API_KEY, model_name="gpt-4")
    response = llm.predict(prompt)
    
    return response

def generate_executive_summary(model_type, statistical_results, dataset_summary, data_quality_results):
    """
    Generates an executive summary for the model validation report using LangChain.
    """
    prompt = f"""
    You are a risk management expert summarizing a model validation report.
    
    Model Type: {model_type}
    
    Statistical Test Results:
    {statistical_results}
    
    Dataset Summary:
    {dataset_summary}
    
    Data Quality Checks:
    {data_quality_results}
    
    Provide a well-structured executive summary that explains the overall findings of the validation exercise, highlights any key issues or concerns, and suggests recommendations for model improvement. Ensure the summary is professional and suitable for a risk management report.
    Do not use headings or bold formatting in your response.
    """
    
    llm = ChatOpenAI(openai_api_key=st.session_state.OPENAI_API_KEY, model_name="gpt-4")
    return llm.predict(prompt)


def generate_model_validation_report(model_type, executive_summary, dataset_summary, data_quality_results, statistical_results, tables_charts, tables_data, table_statuses=None):
    """
    Generates a Word report for model validation, including formatted tables and charts if available.
    """
    doc = Document()
    
    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(f"Model Validation Report for {model_type}")
    run.font.size = Pt(48)
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(executive_summary)
    doc.add_page_break()

    #Table Statuses
    doc.add_heading("Table Statuses", level=1)
    table = doc.add_table(rows=table_statuses.shape[0] + 1, cols=table_statuses.shape[1])

    header_row = table.rows[0].cells
    for j, col_name in enumerate(table_statuses.columns):
        header_row[j].text = col_name

    for i, row_data in enumerate(table_statuses.itertuples(index=False)):
        row_cells = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = str(cell_data)

    format_table_statuses(table)
    
    # Data Summary
    doc.add_heading("Data Summary", level=1)
    doc.add_paragraph(dataset_summary)
    doc.add_page_break()
    
    # Data Quality Analysis
    if data_quality_results:
        doc.add_heading("Data Quality Analysis", level=1)
        for test_name, result in data_quality_results.items():
            doc.add_heading(test_name, level=2)
            doc.add_paragraph(result)

            if test_name in tables_charts:
                for chart_path in tables_charts[test_name]:
                    doc.add_picture(chart_path, width=Inches(5))

            if test_name in tables_data:
                table = doc.add_table(rows=len(tables_data[test_name]), cols=len(tables_data[test_name][0]))
                format_table(table)
                for i, row_data in enumerate(tables_data[test_name]):
                    row_cells = table.rows[i].cells
                    for j, cell_data in enumerate(row_data):
                        row_cells[j].text = str(cell_data)
                        if i == 0:  # Make header row bold
                            row_cells[j].paragraphs[0].runs[0].bold = True
        doc.add_page_break()
    
    # Statistical Analysis
    doc.add_heading("Statistical Analysis", level=1)
    for category, tests in statistical_results.items():
        doc.add_heading(category, level=2)
        for test_name, test_data in tests.items():
            result = test_data.get("analysis", "No analysis available")
            doc.add_heading(test_name, level=3)

            if tables_data is not None and test_name in tables_data:
                table = doc.add_table(rows=len(tables_data[test_name]), cols=len(tables_data[test_name][0]))
                format_table(table)
                for i, row_data in enumerate(tables_data[test_name]):
                    row_cells = table.rows[i].cells
                    for j, cell_data in enumerate(row_data):
                        row_cells[j].text = str(cell_data)
                        if i == 0:
                            row_cells[j].paragraphs[0].runs[0].bold = True
            
                doc.add_paragraph("")

            if tables_charts is not None and test_name in tables_charts:
                for chart_path in tables_charts[test_name]:
                    doc.add_picture(chart_path, width=Inches(5))

            doc.add_paragraph(result)
        
    current_timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_filename = f"Model_Validation_Report_{model_type}_{current_timestamp}.docx"
    doc.save(report_filename)
    
    st.success(f"Report saved as {report_filename}")
    return report_filename



def format_table(table):
    """Applies a table style with borders and bold headers."""
    
    # Apply table style for borders
    table.style = 'Table Grid'  # Word's built-in table style with borders
    
    # Apply bold formatting to the header row
    header_cells = table.rows[0].cells
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True  # Make header row bold



def format_table_statuses(table):
    """Applies a table style with borders, bold headers, reduced font size, and custom column widths."""
    
    # Apply table style for borders
    table.style = 'Table Grid'  # Word's built-in table style with borders
    
    # Apply bold formatting to the header row and reduce font size
    header_cells = table.rows[0].cells
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True  # Make header row bold
                run.font.size = Pt(9)  # Reduce font size to 9 points
    
    # Reduce font size for all other rows
    for row in table.rows[1:]:  # Skip the header row
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)  # Reduce font size to 9 points

    # Set custom column widths
    column_widths = [Inches(1.5), Inches(1.5), Inches(0.5), Inches(2.5)]  # Adjust widths as needed
    for i, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = column_widths[i]  # Set width for each column

